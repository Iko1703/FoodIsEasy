#!/usr/bin/env python3
"""Replace Chapter 2 body in диплом.docx without touching Chapters 1 and 3."""
from pathlib import Path
import shutil
import sys

try:
    from docx import Document
    from docx.oxml.ns import qn
    from docx.text.paragraph import Paragraph
except ImportError:
    print("pip install python-docx", file=sys.stderr)
    sys.exit(1)

BASE = Path(__file__).resolve().parents[1]
DOC_PATH = BASE / "диплом.docx"
CONTENT_PATH = BASE / "_chapter_2_full.txt"
MARKERS_PATH = BASE / "_chapter2_markers.txt"
CH2_MARKER = MARKERS_PATH.read_text(encoding="utf-8").splitlines()[0].strip()
CH3_MARKER = MARKERS_PATH.read_text(encoding="utf-8").splitlines()[1].strip()


def para_text(p: Paragraph) -> str:
    return (p.text or "").strip()


def delete_paragraph(paragraph: Paragraph) -> None:
    p = paragraph._element
    parent = p.getparent()
    if parent is not None:
        parent.remove(p)
        paragraph._p = paragraph._element = None  # type: ignore


def insert_paragraph_after(paragraph: Paragraph, text: str, style: str | None = None) -> Paragraph:
    new_p = paragraph._element.makeelement(qn("w:p"), {})
    paragraph._element.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        try:
            new_para.style = style
        except Exception:
            pass
    new_para.add_run(text)
    return new_para


def main() -> None:
    if not DOC_PATH.exists():
        raise SystemExit(f"Not found: {DOC_PATH}")

    backup = DOC_PATH.with_name(DOC_PATH.stem + "_backup_ch2.docx")
    shutil.copy2(DOC_PATH, backup)
    print(f"Backup: {backup}")

    content = CONTENT_PATH.read_text(encoding="utf-8")
    paragraphs_text = [line.strip() for line in content.splitlines()]

    doc = Document(str(DOC_PATH))
    paras = list(doc.paragraphs)

    ch2_idx = None
    ch3_idx = None
    for i, p in enumerate(paras):
        t = para_text(p)
        if ch2_idx is None and t.startswith(CH2_MARKER):
            ch2_idx = i
        if ch2_idx is not None and i > ch2_idx and t.startswith(CH3_MARKER):
            ch3_idx = i
            break

    if ch2_idx is None:
        raise SystemExit(f"Chapter 2 not found: {CH2_MARKER}")
    if ch3_idx is None:
        raise SystemExit(f"Chapter 3 not found: {CH3_MARKER}")

    # Delete paragraphs between ch2 heading and ch3 heading (exclusive)
    to_delete = paras[ch2_idx + 1 : ch3_idx]
    for p in reversed(to_delete):
        delete_paragraph(p)

    # Refresh anchor after deletions
    paras = list(doc.paragraphs)
    ch2_para = None
    for p in paras:
        if para_text(p).startswith(CH2_MARKER):
            ch2_para = p
            break
    if ch2_para is None:
        raise SystemExit("Chapter 2 anchor lost after delete")

    anchor = ch2_para
    for block in paragraphs_text:
        if not block:
            anchor = insert_paragraph_after(anchor, "")
            continue
        style = None
        if block.startswith("2.") and len(block) > 3 and block[2].isdigit():
            if block.count(".") == 1 or (block.count(".") == 2 and block[4:5].isdigit()):
                style = "Heading 2" if block.count(".") == 1 else "Heading 3"
        anchor = insert_paragraph_after(anchor, block, style)

    out_path = DOC_PATH
    try:
        doc.save(str(out_path))
    except PermissionError:
        out_path = DOC_PATH.with_name(DOC_PATH.stem + "_ch2_updated.docx")
        doc.save(str(out_path))
        print("NOTE: original file locked; saved copy instead.")
    print(f"Updated: {out_path}")


if __name__ == "__main__":
    main()
